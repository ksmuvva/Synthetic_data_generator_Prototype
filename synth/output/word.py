"""
Word output generator using python-docx.

Generates Word documents (.docx) with formatted tables and metadata.
"""

from pathlib import Path
from typing import Optional
import pandas as pd
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    # Create stubs for type hints
    Document = None  # type: ignore

from synth.output.base import OutputGenerator, GeneratorRegistry
from synth.patterns.schema import Schema


class WordGenerator(OutputGenerator):
    """Generate Word document (.docx) output files."""

    def __init__(self):
        """Initialize Word generator."""
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for Word generation. "
                "Install it with: pip install python-docx"
            )

    def generate(
        self,
        data: pd.DataFrame,
        output_path: Path,
        schema: Optional[Schema] = None,
        title: str = "Synthetic Data",
        include_metadata: bool = True,
        **kwargs
    ) -> Path:
        """
        Generate Word document with data table.

        Args:
            data: DataFrame to write
            output_path: Where to save the document
            schema: Optional schema for metadata
            title: Document title
            include_metadata: Include schema metadata if available
            **kwargs: Additional options

        Returns:
            Path to the generated Word document
        """
        # Ensure output path has .docx extension
        output_path = Path(output_path)
        if output_path.suffix not in (".docx", ".doc"):
            output_path = output_path.with_suffix(".docx")

        # Create parent directory if needed
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Create document
        doc = Document()

        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add generation timestamp
        timestamp_para = doc.add_paragraph()
        timestamp_para.text = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        timestamp_run = timestamp_para.runs[0]
        timestamp_run.font.size = Pt(9)
        timestamp_run.font.color.rgb = RGBColor(127, 140, 141)
        timestamp_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add metadata if schema provided
        if schema and include_metadata:
            self._add_metadata_section(doc, schema)

        # Add data table
        self._add_data_table(doc, data)

        # Add footer with record count
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer_run = footer.add_run(f"Total records: {len(data)}")
        footer_run.font.size = Pt(9)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(127, 140, 141)

        # Save document
        doc.save(output_path)

        return output_path

    def _add_metadata_section(self, doc: Document, schema: Schema) -> None:
        """Add metadata section from schema."""
        doc.add_heading("Metadata", level=2)

        # Row count and fields
        doc.add_paragraph(f"Row Count: {schema.row_count}")
        doc.add_paragraph(f"Number of Fields: {len(schema.fields)}")

        # Field details
        if schema.fields:
            doc.add_heading("Fields", level=3)

            for field in schema.fields:
                para = doc.add_paragraph(style="List Bullet")
                field_text = f"{field.name}: {field.type.value}"

                if field.unique:
                    field_text += " [UNIQUE]"
                if not field.nullable:
                    field_text += " [NOT NULL]"

                para.add_run(field_text)

    def _add_data_table(self, doc: Document, data: pd.DataFrame) -> None:
        """Add data table to document."""
        doc.add_heading("Data", level=2)

        # Create table
        table = doc.add_table(rows=1, cols=len(data.columns))
        table.style = "Light Grid Accent 1"

        # Add header row
        header_cells = table.rows[0].cells
        for i, col in enumerate(data.columns):
            cell = header_cells[i]
            cell.text = str(col)
            # Bold header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)

        # Add data rows
        for _, row in data.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                cell = row_cells[i]
                cell.text = str(value)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    def supports_format(self, format_type: str) -> bool:
        """Check if format is supported."""
        return format_type.lower() in ("word", "docx", "doc")


# Register the generator
if PYTHON_DOCX_AVAILABLE:
    GeneratorRegistry.register("word", WordGenerator)
    GeneratorRegistry.register("docx", WordGenerator)
    GeneratorRegistry.register("doc", WordGenerator)
